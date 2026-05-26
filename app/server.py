import json
import os
import re
import shutil
import sys
import traceback
import urllib.error
import urllib.parse
import urllib.request
import uuid
import warnings
import zipfile
from datetime import datetime
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from io import BytesIO

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parent.parent
UPLOADS = ROOT / "uploads"
OUTPUTS = ROOT / "outputs"
STATIC = ROOT / "app"
SAMPLE_DIR = ROOT / "source_zip" / "Meinhardt"


def load_env_file(path):
    if not path.exists():
        return
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


load_env_file(STATIC / ".env")

MIN_FIELD_CONFIDENCE = float(os.environ.get("MIN_FIELD_CONFIDENCE", "0.78"))
GEMINI_TIMEOUT = int(os.environ.get("GEMINI_TIMEOUT", "120"))

for folder in (UPLOADS, OUTPUTS):
    folder.mkdir(exist_ok=True)


class FormField:
    def __init__(self, value, filename=None):
        self.value = value
        self.filename = filename
        self.file = BytesIO(value)

    def __bytes__(self):
        return self.value if isinstance(self.value, bytes) else self.value.encode()


def parse_multipart_form(content_type, body_file):
    boundary = content_type.split("boundary=")[-1].encode().strip()
    parts = {}
    body_file.seek(0)
    body = body_file.read()

    for part in body.split(b"--" + boundary):
        if not part or part == b"--\r\n" or part == b"--":
            continue
        if b"\r\n\r\n" not in part:
            continue
        header_section, content = part.split(b"\r\n\r\n", 1)
        content = content.rstrip(b"\r\n")
        header_text = header_section.decode("utf-8", errors="ignore")

        name_match = re.search(r'name="([^"]+)"', header_text)
        if not name_match:
            continue
        name = name_match.group(1)

        filename_match = re.search(r'filename="([^"]+)"', header_text)
        filename = filename_match.group(1) if filename_match else None

        if name not in parts:
            parts[name] = []
        parts[name].append(FormField(content, filename))

    # Convert single items from list to direct value
    for key in parts:
        if len(parts[key]) == 1:
            parts[key] = parts[key][0]
        # else keep as list for multiple files

    return parts


FIELD_KEYS = (
    "position_title",
    "name",
    "date_of_birth",
    "nationality",
    "country",
    "education",
    "languages",
    "adequacy_tasks",
    "adequacy_skills",
)


KNOWN_TEMPLATE_SCHEMA = {
    "template_name": "Meinhardt CV format",
    "confidence": 0.9,
    "fields": [
        {"field_key": "position_title", "label": "Position Title & No.", "table_index": 0, "row_index": 1, "column_index": 3, "fill_allowed": True, "value_type": "short_text", "notes": "Role/title requested by the client template."},
        {"field_key": "name", "label": "Name of Key Expert", "table_index": 0, "row_index": 2, "column_index": 3, "fill_allowed": True, "value_type": "short_text", "notes": "Candidate full name."},
        {"field_key": "date_of_birth", "label": "Date of Birth", "table_index": 0, "row_index": 3, "column_index": 3, "fill_allowed": True, "value_type": "date_or_text", "notes": "Leave blank unless explicitly present."},
        {"field_key": "nationality", "label": "Nationality", "table_index": 0, "row_index": 4, "column_index": 3, "fill_allowed": True, "value_type": "short_text", "notes": "Leave blank unless stated or strongly implied by explicit resume text."},
        {"field_key": "country", "label": "Country of Citizenship/ Residence", "table_index": 0, "row_index": 5, "column_index": 3, "fill_allowed": True, "value_type": "short_text", "notes": "Country of citizenship or residence."},
        {"field_key": "education", "label": "Education", "table_index": 0, "row_index": 6, "column_index": 3, "fill_allowed": True, "value_type": "long_text", "notes": "Degrees, diplomas, institutions, and years when available."},
        {"field_key": "languages", "label": "Language Skills", "table_index": 0, "row_index": 26, "column_index": 1, "fill_allowed": True, "value_type": "short_text", "notes": "Languages explicitly present in the resume."},
        {"field_key": "adequacy_tasks", "label": "Detailed Tasks Assigned", "table_index": 0, "row_index": 30, "column_index": 1, "fill_allowed": True, "value_type": "long_text", "notes": "Grounded task summary based only on resume evidence."},
        {"field_key": "adequacy_skills", "label": "Assignments and skills", "table_index": 0, "row_index": 30, "column_index": 4, "fill_allowed": True, "value_type": "long_text", "notes": "Project/skills detail grounded in resume evidence."},
    ],
    "employment": {
        "table_index": 0,
        "start_row_index": 9,
        "end_row_index": 23,
        "period_column_index": 1,
        "organization_title_column_index": 2,
        "country_column_index": 4,
        "summary_column_index": 5,
        "fill_allowed": True,
        "notes": "Use one row per relevant employment/project record.",
    },
    "language_ratings": {
        "table_index": 0,
        "row_index": 26,
        "speaking_column_index": 4,
        "reading_column_index": 6,
        "writing_column_index": 7,
        "default_rating": "Good",
        "fill_allowed": True,
    },
}


def safe_name(name):
    cleaned = re.sub(r"[^A-Za-z0-9._ -]+", "_", name or "file")
    return cleaned[:120].strip(" .") or "file"


def extract_docx(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract_pdf(path):
    reader = PdfReader(path)
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def extract_text(path):
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    raise ValueError(f"Unsupported resume type: {suffix}")


def template_snapshot(path):
    doc = Document(path)
    tables = []
    for table_index, table in enumerate(doc.tables):
        rows = []
        for row_index, row in enumerate(table.rows):
            cells = []
            for column_index, cell in enumerate(row.cells):
                text = re.sub(r"\s+", " ", cell.text).strip()
                cells.append({"column_index": column_index, "text": text})
            rows.append({"row_index": row_index, "cells": cells})
        tables.append({"table_index": table_index, "rows": rows})
    paragraphs = [re.sub(r"\s+", " ", p.text).strip() for p in doc.paragraphs if p.text.strip()]
    return {"paragraphs": paragraphs, "tables": tables}


def field_schema():
    return {
        "type": "object",
        "additionalProperties": False,
        "required": ["value", "confidence", "source_quote", "fill_allowed"],
        "properties": {
            "value": {"type": "string"},
            "confidence": {"type": "number", "minimum": 0, "maximum": 1},
            "source_quote": {"type": "string"},
            "fill_allowed": {"type": "boolean"},
        },
    }


TEMPLATE_JSON_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "required": ["template_name", "confidence", "fields", "employment", "language_ratings"],
    "properties": {
        "template_name": {"type": "string"},
        "confidence": {"type": "number", "minimum": 0, "maximum": 1},
        "fields": {
            "type": "array",
            "items": {
                "type": "object",
                "additionalProperties": False,
                "required": ["field_key", "label", "table_index", "row_index", "column_index", "fill_allowed", "value_type", "notes"],
                "properties": {
                    "field_key": {"type": "string", "enum": list(FIELD_KEYS)},
                    "label": {"type": "string"},
                    "table_index": {"type": "integer", "minimum": 0},
                    "row_index": {"type": "integer", "minimum": 0},
                    "column_index": {"type": "integer", "minimum": 0},
                    "fill_allowed": {"type": "boolean"},
                    "value_type": {"type": "string", "enum": ["short_text", "long_text", "date_or_text"]},
                    "notes": {"type": "string"},
                },
            },
        },
        "employment": {
            "type": "object",
            "additionalProperties": False,
            "required": [
                "table_index",
                "start_row_index",
                "end_row_index",
                "period_column_index",
                "organization_title_column_index",
                "country_column_index",
                "summary_column_index",
                "fill_allowed",
                "notes",
            ],
            "properties": {
                "table_index": {"type": "integer", "minimum": 0},
                "start_row_index": {"type": "integer", "minimum": 0},
                "end_row_index": {"type": "integer", "minimum": 0},
                "period_column_index": {"type": "integer", "minimum": 0},
                "organization_title_column_index": {"type": "integer", "minimum": 0},
                "country_column_index": {"type": "integer", "minimum": 0},
                "summary_column_index": {"type": "integer", "minimum": 0},
                "fill_allowed": {"type": "boolean"},
                "notes": {"type": "string"},
            },
        },
        "language_ratings": {
            "type": "object",
            "additionalProperties": False,
            "required": [
                "table_index",
                "row_index",
                "speaking_column_index",
                "reading_column_index",
                "writing_column_index",
                "default_rating",
                "fill_allowed",
            ],
            "properties": {
                "table_index": {"type": "integer", "minimum": 0},
                "row_index": {"type": "integer", "minimum": 0},
                "speaking_column_index": {"type": "integer", "minimum": 0},
                "reading_column_index": {"type": "integer", "minimum": 0},
                "writing_column_index": {"type": "integer", "minimum": 0},
                "default_rating": {"type": "string"},
                "fill_allowed": {"type": "boolean"},
            },
        },
    },
}


CANDIDATE_JSON_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "required": [
        "name",
        "position_title",
        "date_of_birth",
        "nationality",
        "country",
        "education",
        "languages",
        "email",
        "phone",
        "adequacy_tasks",
        "adequacy_skills",
        "employment",
        "review_notes",
    ],
    "properties": {
        "name": field_schema(),
        "position_title": field_schema(),
        "date_of_birth": field_schema(),
        "nationality": field_schema(),
        "country": field_schema(),
        "education": field_schema(),
        "languages": field_schema(),
        "email": field_schema(),
        "phone": field_schema(),
        "adequacy_tasks": field_schema(),
        "adequacy_skills": field_schema(),
        "employment": {
            "type": "array",
            "items": {
                "type": "object",
                "additionalProperties": False,
                "required": [
                    "period",
                    "organization_title",
                    "country",
                    "summary",
                    "confidence",
                    "source_quote",
                    "fill_allowed",
                ],
                "properties": {
                    "period": {"type": "string"},
                    "organization_title": {"type": "string"},
                    "country": {"type": "string"},
                    "summary": {"type": "string"},
                    "confidence": {"type": "number", "minimum": 0, "maximum": 1},
                    "source_quote": {"type": "string"},
                    "fill_allowed": {"type": "boolean"},
                },
            },
        },
        "review_notes": {"type": "array", "items": {"type": "string"}},
    },
}


def extract_gemini_response_text(data):
    chunks = []
    for candidate in data.get("candidates", []):
        content = candidate.get("content") or {}
        for part in content.get("parts", []):
            if part.get("text"):
                chunks.append(part["text"])
    return "\n".join(chunks)


def gemini_response_schema(schema):
    unsupported = {"additionalProperties", "minimum", "maximum"}
    if isinstance(schema, dict):
        return {key: gemini_response_schema(value) for key, value in schema.items() if key not in unsupported}
    if isinstance(schema, list):
        return [gemini_response_schema(item) for item in schema]
    return schema


def call_gemini_json(prompt, schema_name, schema):
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        return None
    instruction = (
        "You convert resumes into client CV templates. "
        "Return only facts supported by the supplied source text. "
        "Do not invent missing values. Use empty strings and low confidence when evidence is absent.\n\n"
        f"Return JSON for schema '{schema_name}'.\n\n"
    )
    payload = {
        "contents": [
            {
                "role": "user",
                "parts": [{"text": instruction + prompt}],
            }
        ],
        "generationConfig": {
            "responseMimeType": "application/json",
            "responseSchema": gemini_response_schema(schema),
        },
    }
    model = os.environ.get("GEMINI_MODEL", "gemini-3.5-flash")
    request = urllib.request.Request(
        f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent",
        data=json.dumps(payload).encode("utf-8"),
        headers={"x-goog-api-key": api_key, "Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=GEMINI_TIMEOUT) as response:
            data = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = exc.reason or f"HTTP {exc.code}"
        try:
            body = json.loads(exc.read().decode("utf-8"))
            detail = body.get("error", {}).get("message") or detail
        except Exception:
            pass
        raise RuntimeError(f"Gemini request failed ({exc.code}): {detail}") from exc
    return json.loads(extract_gemini_response_text(data))


def first_match(text, patterns, default=""):
    for pattern in patterns:
        match = re.search(pattern, text, re.I | re.M)
        if match:
            value = match.group(1).strip(" :,-\n\t")
            if value:
                return re.sub(r"\s+", " ", value)
    return default


def detect_name(text, filename):
    for line in text.splitlines()[:18]:
        line = re.sub(r"\s+", " ", line).strip()
        line = re.sub(r"^(name|candidate name|applicant name)\s*[:\-]\s*", "", line, flags=re.I).strip()
        if not line or "curriculum" in line.lower() or "resume" in line.lower():
            continue
        if re.search(r"@|www\.|linkedin|mobile|phone|\d{4,}", line, re.I):
            continue
        words = [w for w in re.split(r"\s+", line) if w]
        if 2 <= len(words) <= 5 and sum(ch.isalpha() for ch in line) > 5:
            return line.title()
    return Path(filename).stem.replace("_", " ").replace("-", " ").title()


def split_sections(text):
    lines = [re.sub(r"\s+", " ", l).strip() for l in text.splitlines() if l.strip()]
    return lines


def extract_education(lines):
    keywords = r"(b\.?tech|bachelor|diploma|degree|iti|polytechnic|m\.?tech|university|college|examinations board)"
    found = [l for l in lines if re.search(keywords, l, re.I) and len(l) < 220]
    found.sort(key=lambda line: 0 if re.search(r"\b(from|university|college|board|school|institute)\b", line, re.I) else 1)
    return "; ".join(found[:4])


def extract_languages(text):
    explicit = first_match(text, [r"languages?\s*[:\-]\s*(.+)", r"language skills?\s*[:\-]\s*(.+)"])
    if explicit:
        return explicit[:180]
    langs = []
    for lang in ("English", "Hindi", "Gujarati", "Marathi", "Urdu", "Tamil", "Telugu", "Kannada"):
        if re.search(rf"\b{lang}\b", text, re.I):
            langs.append(lang)
    return ", ".join(dict.fromkeys(langs)) or "English"


def extract_experience(lines):
    rows = []
    date_re = r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)?\.?\s*\d{4}|\d{1,2}[/-]\d{4}|\d{4})\s*(?:-|to|–|—)\s*((?:Present|Till Date|Current)|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)?\.?\s*\d{4}|\d{1,2}[/-]\d{4}|\d{4})"
    for i, line in enumerate(lines):
        if re.search(date_re, line, re.I):
            context = " ".join(lines[i : min(i + 4, len(lines))])
            rows.append(
                {
                    "period": re.search(date_re, line, re.I).group(0),
                    "organization_title": context[:280],
                    "country": first_match(context, [r"\b(India|UAE|Qatar|Saudi Arabia|Oman|Kuwait|Bahrain)\b"], "India"),
                    "summary": context[:420],
                }
            )
        if len(rows) >= 12:
            break
    if not rows:
        summary_lines = [l for l in lines if re.search(r"engineer|manager|survey|cadd|scada|mep|project|site|supervisor", l, re.I)]
        rows.append(
            {
                "period": "",
                "organization_title": summary_lines[0][:240] if summary_lines else "",
                "country": "India",
                "summary": "; ".join(summary_lines[:4])[:420],
            }
        )
    return rows


def heuristic_parse(text, filename):
    lines = split_sections(text)
    email = first_match(text, [r"([A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,})"])
    phone = first_match(text, [r"(\+?\d[\d\s().-]{8,}\d)"])
    designation = first_match(
        text,
        [
            r"(?:position|designation|post applied for|title)\s*[:\-]\s*(.+)",
            r"\b(Sr\.?\s+[^,\n]{4,60}|Senior\s+[^,\n]{4,60}|[A-Za-z& ]+\s+Engineer|Surveyor|CADD\s+Expert|SCADA\s+Expert)\b",
        ],
    )
    return {
        "name": detect_name(text, filename),
        "position_title": designation,
        "date_of_birth": first_match(text, [r"(?:date of birth|dob|d\.o\.b)\s*[:\-]?\s*([0-9A-Za-z ,/.-]{6,32})"]),
        "nationality": first_match(text, [r"nationality\s*[:\-]\s*([A-Za-z ]{3,40})"], "Indian"),
        "country": first_match(text, [r"(?:country of residence|country of citizenship|residence country)\s*[:\-]\s*([A-Za-z ,.-]{3,80})"], "India"),
        "education": extract_education(lines),
        "languages": extract_languages(text),
        "email": email,
        "phone": phone,
        "employment": extract_experience(lines),
        "adequacy_tasks": "To be aligned with the client's assignment scope.",
        "adequacy_skills": "; ".join([l for l in lines if re.search(r"skills?|software|autocad|revit|scada|mep|survey", l, re.I)][:6])[:700],
    }


def analyze_template(template_path):
    snapshot = template_snapshot(template_path)
    prompt = (
        "Analyze this DOCX client CV template snapshot and return the fill plan JSON. "
        "Identify only cells that should be filled with candidate data. "
        "Use field keys from the schema enum only. Employment must describe the repeatable table row area. "
        "If the template matches the Meinhardt format, map the general fields, employment rows, language row, "
        "and adequacy section exactly.\n\n"
        f"Template filename: {template_path.name}\nTemplate snapshot:\n{json.dumps(snapshot)[:24000]}"
    )
    fallback = json.loads(json.dumps(KNOWN_TEMPLATE_SCHEMA))
    try:
        ai_schema = call_gemini_json(prompt, "cv_template_fill_plan", TEMPLATE_JSON_SCHEMA)
        if ai_schema:
            ai_schema["ai_used"] = True
            return ai_schema
        fallback["ai_used"] = False
        fallback["ai_error"] = "GEMINI_API_KEY is not set; used built-in Meinhardt template map."
        return fallback
    except Exception as exc:
        fallback["ai_used"] = False
        fallback["ai_error"] = str(exc)
        return fallback


def trusted_field(field):
    if not isinstance(field, dict):
        return False
    value = str(field.get("value") or "").strip()
    if not value:
        return False
    if not field.get("fill_allowed"):
        return False
    try:
        confidence = float(field.get("confidence") or 0)
    except (TypeError, ValueError):
        confidence = 0
    return confidence >= MIN_FIELD_CONFIDENCE


def trusted_record(record):
    if not isinstance(record, dict) or not record.get("fill_allowed"):
        return False
    if not any(str(record.get(k) or "").strip() for k in ("period", "organization_title", "summary")):
        return False
    try:
        confidence = float(record.get("confidence") or 0)
    except (TypeError, ValueError):
        confidence = 0
    return confidence >= MIN_FIELD_CONFIDENCE


def flatten_ai_candidate(ai_data):
    flattened = {}
    held_back = []
    for key in FIELD_KEYS + ("email", "phone"):
        field = ai_data.get(key, {}) if isinstance(ai_data, dict) else {}
        if trusted_field(field):
            flattened[key] = str(field.get("value") or "").strip()
        else:
            flattened[key] = ""
            if isinstance(field, dict) and field.get("value"):
                held_back.append({"field": key, "value": field.get("value"), "confidence": field.get("confidence"), "reason": "below threshold or fill not allowed"})
    employment = []
    for record in ai_data.get("employment", []) if isinstance(ai_data, dict) else []:
        if trusted_record(record):
            employment.append(
                {
                    "period": str(record.get("period") or "").strip(),
                    "organization_title": str(record.get("organization_title") or "").strip(),
                    "country": str(record.get("country") or "").strip(),
                    "summary": str(record.get("summary") or "").strip(),
                }
            )
        elif isinstance(record, dict) and any(record.get(k) for k in ("period", "organization_title", "summary")):
            held_back.append({"field": "employment", "value": record, "confidence": record.get("confidence"), "reason": "below threshold or fill not allowed"})
    flattened["employment"] = employment
    flattened["ai_raw"] = ai_data
    flattened["held_back"] = held_back
    flattened["review_notes"] = ai_data.get("review_notes", []) if isinstance(ai_data, dict) else []
    return flattened


def call_gemini_resume_parser(text, filename, template_schema):
    prompt = (
        "Extract this resume into the provided client template fill fields. "
        "For every value, include a short source quote copied from the resume text. "
        "Set fill_allowed false for values that are inferred, unsupported, ambiguous, or not intended for the template. "
        "Use confidence below 0.78 when the value should be reviewed instead of filled. "
        "For adequacy fields, summarize only project/tasks/skills actually supported by resume evidence. "
        "Leave fields empty when the resume does not clearly provide the information.\n\n"
        f"Filename: {filename}\n"
        f"Template fill plan:\n{json.dumps(template_schema)[:16000]}\n\n"
        f"Resume text:\n{text[:42000]}"
    )
    return call_gemini_json(prompt, "candidate_cv_extraction", CANDIDATE_JSON_SCHEMA)


def parse_resume(path, template_schema):
    text = extract_text(path)
    fallback = heuristic_parse(text, path.name)
    try:
        ai = call_gemini_resume_parser(text, path.name, template_schema)
        if ai:
            parsed = flatten_ai_candidate(ai)
            for key, value in fallback.items():
                if key not in parsed or parsed[key] in ("", [], None):
                    parsed[key] = value if not os.environ.get("GEMINI_API_KEY") else parsed.get(key, "")
            parsed["ai_used"] = True
        else:
            parsed = fallback
            parsed["ai_used"] = False
            parsed["ai_error"] = "GEMINI_API_KEY is not set; used local rule-based extraction."
    except Exception as exc:
        parsed = fallback
        parsed["ai_used"] = False
        parsed["ai_error"] = str(exc)
    parsed["source_file"] = path.name
    parsed["raw_text_preview"] = text[:1200]
    return parsed


def set_cell(cell, value):
    cell.text = str(value or "")
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = None


def table_cell(doc, table_index, row_index, column_index):
    if table_index >= len(doc.tables):
        return None
    table = doc.tables[table_index]
    if row_index >= len(table.rows) or column_index >= len(table.rows[row_index].cells):
        return None
    return table.cell(row_index, column_index)


def fill_docx_template(template_path, data, template_schema, output_path):
    shutil.copyfile(template_path, output_path)
    doc = Document(output_path)
    if not doc.tables:
        raise ValueError("Template must contain at least one table.")
    audit = {"filled_fields": [], "blank_fields": [], "held_back": data.get("held_back", [])}

    for field in template_schema.get("fields", []):
        key = field.get("field_key")
        value = data.get(key)
        if not field.get("fill_allowed") or not value:
            audit["blank_fields"].append({"field": key, "reason": "missing, low confidence, or not fillable"})
            continue
        cell = table_cell(doc, field.get("table_index", 0), field.get("row_index", 0), field.get("column_index", 0))
        if cell is None:
            audit["blank_fields"].append({"field": key, "reason": "target cell not found"})
            continue
        set_cell(cell, value)
        audit["filled_fields"].append(key)

    employment = data.get("employment") or []
    employment_schema = template_schema.get("employment", {})
    if employment_schema.get("fill_allowed", True):
        start = employment_schema.get("start_row_index", 9)
        end = employment_schema.get("end_row_index", 23)
        max_rows = max(0, end - start + 1)
        for offset, job in enumerate(employment[:max_rows], start=start):
            targets = (
                ("period", employment_schema.get("period_column_index", 1)),
                ("organization_title", employment_schema.get("organization_title_column_index", 2)),
                ("country", employment_schema.get("country_column_index", 4)),
                ("summary", employment_schema.get("summary_column_index", 5)),
            )
            for key, column_index in targets:
                cell = table_cell(doc, employment_schema.get("table_index", 0), offset, column_index)
                if cell is not None:
                    set_cell(cell, job.get(key, ""))
            audit["filled_fields"].append(f"employment[{offset - start}]")
    elif employment:
        audit["blank_fields"].append({"field": "employment", "reason": "template schema does not allow employment fill"})

    language_schema = template_schema.get("language_ratings", {})
    if data.get("languages") and language_schema.get("fill_allowed", True):
        for key in ("speaking_column_index", "reading_column_index", "writing_column_index"):
            cell = table_cell(doc, language_schema.get("table_index", 0), language_schema.get("row_index", 26), language_schema.get(key, 0))
            if cell is not None:
                set_cell(cell, language_schema.get("default_rating") or "Good")
    doc.save(output_path)
    return audit


def make_job(template_item, resume_items):
    job_id = datetime.now().strftime("%Y%m%d-%H%M%S") + "-" + uuid.uuid4().hex[:6]
    upload_dir = UPLOADS / job_id
    out_dir = OUTPUTS / job_id
    upload_dir.mkdir(parents=True)
    out_dir.mkdir(parents=True)

    template_path = upload_dir / safe_name(template_item.filename)
    with open(template_path, "wb") as handle:
        handle.write(template_item.file.read())

    template_schema = analyze_template(template_path)
    (out_dir / "template-schema.json").write_text(json.dumps(template_schema, indent=2), encoding="utf-8")

    results = []
    for item in resume_items:
        resume_path = upload_dir / safe_name(item.filename)
        with open(resume_path, "wb") as handle:
            handle.write(item.file.read())
        parsed = parse_resume(resume_path, template_schema)
        output_name = f"{safe_name(parsed.get('name') or resume_path.stem)} - client format.docx"
        output_path = out_dir / output_name
        fill_audit = fill_docx_template(template_path, parsed, template_schema, output_path)
        candidate_json_name = f"{safe_name(resume_path.stem)} - extraction.json"
        (out_dir / candidate_json_name).write_text(json.dumps(parsed, indent=2), encoding="utf-8")
        results.append(
            {
                "source": item.filename,
                "name": parsed.get("name"),
                "position": parsed.get("position_title"),
                "ai_used": parsed.get("ai_used"),
                "ai_error": parsed.get("ai_error", ""),
                "needs_review": bool(fill_audit.get("held_back") or parsed.get("review_notes")),
                "download": f"/outputs/{job_id}/{output_name}",
                "extraction": f"/outputs/{job_id}/{candidate_json_name}",
                "fill_audit": fill_audit,
                "data": parsed,
            }
        )

    manifest_path = out_dir / "extraction-results.json"
    manifest_path.write_text(json.dumps(results, indent=2), encoding="utf-8")
    zip_path = OUTPUTS / f"{job_id}.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for file in out_dir.iterdir():
            archive.write(file, file.name)
    return {"job_id": job_id, "results": results, "zip": f"/outputs/{zip_path.name}"}


class Handler(BaseHTTPRequestHandler):
    def send_json(self, payload, status=200):
        body = json.dumps(payload).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_GET(self):
        path = urllib.parse.unquote(self.path.split("?", 1)[0])
        if path == "/":
            path = "/index.html"
        if path.startswith("/outputs/"):
            file_path = OUTPUTS / path.replace("/outputs/", "", 1)
        else:
            file_path = STATIC / path.lstrip("/")
        if not file_path.exists() or not file_path.is_file():
            self.send_error(404)
            return
        content_type = "text/html"
        if file_path.suffix == ".css":
            content_type = "text/css"
        elif file_path.suffix == ".js":
            content_type = "application/javascript"
        elif file_path.suffix == ".docx":
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif file_path.suffix == ".zip":
            content_type = "application/zip"
        elif file_path.suffix == ".json":
            content_type = "application/json"
        data = file_path.read_bytes()
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def do_POST(self):
        if self.path != "/api/process":
            self.send_error(404)
            return
        try:
            content_length = int(self.headers.get("Content-Length", 0))
            body_file = BytesIO(self.rfile.read(content_length))
            content_type = self.headers.get("Content-Type", "")
            form = parse_multipart_form(content_type, body_file)

            template = form.get("template") if "template" in form else None
            resumes = form.get("resumes") if "resumes" in form else []
            if not isinstance(resumes, list):
                resumes = [resumes]
            resumes = [item for item in resumes if getattr(item, "filename", "")]
            if template is None or not getattr(template, "filename", ""):
                self.send_json({"error": "Upload a client DOCX template first."}, 400)
                return
            if not resumes:
                self.send_json({"error": "Upload at least one resume."}, 400)
                return
            self.send_json(make_job(template, resumes))
        except Exception:
            traceback.print_exc()
            self.send_json({"error": traceback.format_exc()}, 500)


def main():
    host = "127.0.0.1"
    port = int(os.environ.get("PORT", "8765"))
    print(f"CV automation app running at http://{host}:{port}")
    ThreadingHTTPServer((host, port), Handler).serve_forever()


if __name__ == "__main__":
    main()
